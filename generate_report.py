import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = docx.Document()

# Set styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

heading_style = doc.styles['Heading 1']
h_font = heading_style.font
h_font.name = 'Times New Roman'
h_font.size = Pt(14)
h_font.bold = True
h_font.color.rgb = docx.shared.RGBColor(0, 0, 0)

sub_heading_style = doc.styles['Heading 2']
sh_font = sub_heading_style.font
sh_font.name = 'Times New Roman'
sh_font.size = Pt(12)
sh_font.bold = True
sh_font.color.rgb = docx.shared.RGBColor(0, 0, 0)

# Function to add paragraph with 1.5 spacing
def add_para(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if bold:
        p.add_run(text).bold = True
    else:
        p.add_run(text)
    return p

# 1. Cover Page
p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("\n\n\n\nPROJECT REPORT ON\n")
run.font.size = Pt(16)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("ToolForge — Agentic AI Workflow Platform\n")
run.font.size = Pt(20)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("\nSubmitted in partial fulfillment of the requirements for the degree of\n[Degree / Course]\nin\n[Branch]\n\n")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("Submitted by:\n")
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run("Bada Ravi Kiran Reddy (Roll No: __________)\n")
p.add_run("V. Tanish (Roll No: __________)\n")
p.add_run("Kandunuri Eekshith Sai (Roll No: __________)\n")

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("\n\nUnder the Guidance of:\n[Guide Name / Teacher Name]\n")
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = p.add_run("\n\n\n[College Name]\n[Submission Date]\n")
run.font.size = Pt(16)
run.bold = True

# No page break
doc.add_paragraph("\n\n")

# 2. Index
doc.add_heading("Table of Contents", level=1)
index_text = """1. Problem Statement .............................................................. 2
2. Jira Timeline .................................................................. 2
3. Requirements (SRS) ............................................................. 3
4. Design (SDD) ................................................................... 4
5. Technical Landscape ............................................................ 5
6. Output ......................................................................... 7
7. Conclusion & Future Scope ...................................................... 8"""
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.add_run(index_text)

# Content sections
def add_section_heading(title):
    doc.add_paragraph("\n") # ONE extra blank line before new section
    doc.add_heading(title, level=1)
    
add_section_heading("1. Problem Statement")
add_para("Enterprise technical tasks such as code reviews, SQL generation, and workflow planning are often fragmented across multiple siloed tools. This lack of integration leads to context loss, significant operational inefficiency, and reduced productivity among technical teams. The conventional process requires users to manually switch between development environments, databases, and research tools.")
add_para("ToolForge is proposed to solve this by providing a unified, secure command center featuring specialized agents with built-in LangGraph-powered ReAct (Reasoning & Acting) logic. By integrating directly into a high-performance MERN architecture with enterprise-grade authentication, ToolForge aims to automate and orchestrate complex enterprise workflows into autonomous AI tasks.")

add_section_heading("2. Jira Timeline")
add_para("The project was executed using an agile, sprint-based approach. The timeline is divided into iterations, tracking progress from design to deployment.")
add_para("Sprint 1 (Week 1-2): Requirement Analysis & Design")
add_para("• Task 1: Finalize functional requirements and use cases.\n• Task 2: Design system architecture and data flow diagrams.\n• Task 3: Create initial UI/UX mockups for the dashboard.")
add_para("Sprint 2 (Week 3-4): Backend & Database Development")
add_para("• Task 1: Set up Node.js Express server and MongoDB Atlas.\n• Task 2: Implement Passport.js authentication (Local + Google OAuth).\n• Task 3: Develop core API routing structures.")
add_para("Sprint 3 (Week 5-6): Frontend Development & UI/UX")
add_para("• Task 1: Build React 19 application with Vite.\n• Task 2: Implement glassmorphism UI using Tailwind/Custom CSS.\n• Task 3: Connect frontend context with backend APIs.")
add_para("Sprint 4 (Week 7-8): Agentic AI Integration")
add_para("• Task 1: Integrate LangGraph.js ReAct loop and Groq Llama 3.3.\n• Task 2: Develop Web Research, MongoDB Query, and Code Review agents.\n• Task 3: Develop Workflow Planner, Prompt Engineering, and API Integration agents.")
add_para("Sprint 5 (Week 9-10): DevOps & Deployment")
add_para("• Task 1: Containerize application using Docker & docker-compose.\n• Task 2: Set up Jenkins CI/CD pipeline.\n• Task 3: Final deployment to Render (Backend) and Vercel (Frontend).")

add_section_heading("3. Requirements (SRS)")
doc.add_heading("3.1 Functional Requirements", level=2)
add_para("• User Authentication: System must support secure login using local credentials and Google OAuth 2.0.\n• Agent Interaction: Users can interact with 6 specialized AI agents (Web Research, MongoDB Query, Code Review, Workflow Planner, Prompt Engineering, API Integration).\n• Autonomous Execution: System must utilize LangGraph to reason, execute tools, and trace workflows autonomously.\n• Dashboard: System must display an intuitive, reactive dashboard mapping AI outputs dynamically.")

doc.add_heading("3.2 Non-Functional Requirements", level=2)
add_para("• Performance: The React frontend must render smoothly with minimal load times via Vite optimization.\n• Security: System must utilize JWT for secure sessions and HTTPS/SSL for deployment endpoints.\n• Portability: The application must be fully containerized using Docker, allowing it to run uniformly across diverse environments.\n• Maintainability: Code must be modular, adhering to MERN stack best practices, with version control managed through Git.")

add_section_heading("4. Design (SDD)")
doc.add_heading("4.1 System Architecture", level=2)
add_para("ToolForge is built on a Pure MERN (MongoDB, Express, React, Node.js) architecture with an embedded Agentic reasoning loop.")
add_para("• Frontend: React UI sends API requests to the Backend.\n• Backend: Node.js Express server handles routing and passes relevant requests to the Agent Router.\n• Agentic Loop: Inside the application layer, LangGraph.js runs a ReAct (Reasoning + Acting) loop utilizing Groq's Llama 3.3 model.\n• External Services: The agents interact with external utilities like MongoDB Atlas (via MongoDB Tool) and Tavily Web API (via Search Tool) to perform specialized tasks.")

doc.add_heading("4.2 Data Flow Diagram (DFD)", level=2)
add_para("1. User provides a prompt or request via the React Dashboard.\n2. The Dashboard component transmits an HTTP request containing the prompt and authentication headers to the Express backend.\n3. The Express router validates the JWT and routes the request to the LangGraph logic.")
add_para("4. The LLM (Llama 3.3) analyzes the prompt, generating a LangGraph-powered ReAct (Reasoning & Acting) loop.")
add_para("5. The LLM dynamically triggers required Tools (e.g., executing MQL on MongoDB).\n6. Final computed output is synthesized and returned through the Express backend to the React UI, updating the application state.")

doc.add_heading("4.3 Database Design", level=2)
add_para("MongoDB is utilized as the primary database, optimized for document-oriented flexibility. Collections manage User records (credentials, OAuth IDs), session data, and agent execution logs. Mongoose schemas are employed in the backend to ensure strict data validation and structure enforcement.")

add_section_heading("5. Technical Landscape")
doc.add_heading("5.1 Git", level=2)
add_para("Git is utilized for version control, strictly following a Feature-Branch workflow. The 'main' branch mirrors production and is locked for direct commits, while 'develop' acts as the integration branch. This allows safe, isolated development of complex AI agent features. This version control strategy is essential for tracking progress and collaborating efficiently.")

doc.add_heading("5.2 Docker", level=2)
add_para("Docker is employed to containerize the entire stack, providing a 'Run Anywhere' portability layer. A Dockerfile defines the Node.js application image, while 'docker-compose.yml' orchestrates the multi-container setup (including a MongoDB 7.0 container), simplifying local development and CI/CD operations. This eliminates the 'it works on my machine' problem.")

doc.add_heading("5.3 Jenkins", level=2)
add_para("Jenkins drives the Continuous Integration and Continuous Deployment (CI/CD) pipeline. A scripted 'Jenkinsfile' automates essential stages: source checkout, Docker build/deploy, and health checks, ensuring all code modifications are systematically validated before being served. This automates the release process securely.")

doc.add_heading("5.4 MERN Stack", level=2)
add_para("The MERN stack provides the backbone of the application. React.js and Vite deliver a highly responsive, modern client-side experience. Node.js and Express form a non-blocking, scalable backend suited for handling long-polling agentic requests. MongoDB Atlas provides a resilient, cloud-native storage layer essential for rapid iterative querying. These technologies were chosen because they utilize JavaScript throughout the stack, increasing development speed.")

doc.add_heading("5.5 Other Libraries / External APIs", level=2)
add_para("• Groq / Llama 3.3: Provides the ultra-fast LLM inference essential for the core AI reasoning engine.\n• LangGraph.js: Enables complex, stateful multi-agent workflows and ReAct-based execution.\n• Passport.js & JWT: Standardizes and secures the authentication flow.\n• Tavily: An optimized web search API that empowers the Web Research agent with real-time data access.")

add_section_heading("6. Output")
add_para("The deployed system delivers a high-performance web interface available at toolforge-lyart.vercel.app. Upon navigation, unauthenticated users view the landing page, detailing platform capabilities. After logging in (via Google OAuth or local credentials), users are redirected to the Dashboard.")
add_para("Step-by-step Execution:\n1. The user inputs a complex task (e.g., 'Generate a MongoDB query to find all active users').\n2. The request is submitted; the UI displays a processing state.\n3. In the backend, the LangGraph engine intercepts the request, deciding which specialized agent to deploy.\n4. The agent executes the underlying reasoning, optionally interacting with real databases or searching the web.\n5. A final processed response, containing optimized code or research findings, is securely returned to the user's dashboard.")
doc.add_paragraph("[Insert Dashboard Screenshot Here]", style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Figure 1: ToolForge Agentic Dashboard Layout", style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_section_heading("7. Conclusion & Future Scope")
add_para("ToolForge successfully demonstrates the integration of autonomous Agentic AI loops directly into a production-ready MERN application. By wrapping LLM reasoning with strict programmatic tools and maintaining a robust DevOps lifecycle (Docker, Jenkins), the project elevates AI from a simple chatbot to a deterministic workflow engine. The platform successfully bridges the gap between complex enterprise operations and intuitive AI execution.")
add_para("Future Scope:\n• Expansion of the agent library to support automated cloud infrastructure provisioning (Terraform/AWS).\n• Implementation of collaborative multiplayer dashboards where multiple users can interact with the same agentic trace.\n• Enhancing the observability pipelines using LangSmith for more granular AI debugging and performance tracking.")

# Signatures Space
doc.add_paragraph("\n\n\n\n\n")
p = doc.add_paragraph()
# We will create a pseudo table or tabs for signatures
p.add_run("______________________                  ______________________                  ______________________\n")
p.add_run("Bada Ravi Kiran Reddy                         V. Tanish                                    Kandunuri Eekshith Sai")
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.save('/Users/ravikiranreddybada/Documents/toolforge/ToolForge_Project_Report.docx')
